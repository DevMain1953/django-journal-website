from django.conf import settings
from docx import Document
from typing import Dict


class ArticleManager:
    def __init__(self):
        self.__docx_document_parser = Document
    

    def get_name_description_and_authors_of_article_in_russian_from_file(self, file_name: str) -> Dict[str, str]:
        result = {}
        number_of_paragraph = 0
        path_to_file = settings.MEDIA_ROOT + file_name
        document = self.__docx_document_parser(path_to_file)
        for paragraph in document.paragraphs:
            if paragraph.style.name == "Main Title":
                number_of_paragraph += 1
                if number_of_paragraph == 2:
                    result["name"] = paragraph.text
                if number_of_paragraph == 3:
                    result["short_description"] = paragraph.text
        number_of_paragraph = 0
        for paragraph in document.paragraphs:
            if paragraph.style.name == "University Name":
                number_of_paragraph += 1
                if number_of_paragraph == 1:
                    result["authors"] = paragraph.text
        return result
    

    def get_name_description_and_authors_of_article_in_english_from_file(self, file_name: str) -> Dict[str, str]:
        result = {}
        number_of_paragraph = 0
        path_to_file = settings.MEDIA_ROOT + file_name
        document = self.__docx_document_parser(path_to_file)
        for paragraph in document.paragraphs:
            if paragraph.style.name == "Main Title":
                number_of_paragraph += 1
                if number_of_paragraph == 4:
                    result["name"] = paragraph.text
                if number_of_paragraph == 5:
                    result["short_description"] = paragraph.text
        number_of_paragraph = 0
        for paragraph in document.paragraphs:
            if paragraph.style.name == "University Name":
                number_of_paragraph += 1
                if number_of_paragraph == 3:
                    result["authors"] = paragraph.text
        return result
